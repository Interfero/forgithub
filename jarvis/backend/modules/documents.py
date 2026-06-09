"""
Режим «Бухгалтер + Юрист»: контрагенты (SQLite), выписки, генерация .docx/.xlsx.
"""

from __future__ import annotations

import json
import re
import sqlite3
import uuid
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

from modules.app_paths import user_data_dir

DATA_DIR = user_data_dir()
DB_PATH = DATA_DIR / "accountant.db"
GENERATED_DIR = DATA_DIR / "generated"
UPLOADS_DIR = DATA_DIR / "uploads" / "accountant"

# Последняя аналитика выписки для подмешивания в промпт
_last_statement_summary: str = ""


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def init_db() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS counterparties (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL DEFAULT '',
                inn TEXT,
                kpp TEXT,
                ogrn TEXT,
                account TEXT,
                bik TEXT,
                raw_text TEXT,
                created_at TEXT NOT NULL
            )
            """
        )
        conn.execute(
            "CREATE UNIQUE INDEX IF NOT EXISTS idx_cp_inn ON counterparties(inn) WHERE inn IS NOT NULL"
        )


def parse_requisites(text: str) -> dict[str, str]:
    """Извлечение реквизитов из текста сообщения пользователя."""
    t = text.replace("\n", " ")
    out: dict[str, str] = {"name": ""}

    inn = re.search(r"\bИНН[:\s]*(\d{10}|\d{12})\b", text, re.I) or re.search(
        r"\b(\d{10}|\d{12})\b", t
    )
    if inn:
        out["inn"] = inn.group(1)

    kpp = re.search(r"\bКПП[:\s]*(\d{9})\b", text, re.I)
    if kpp:
        out["kpp"] = kpp.group(1)

    ogrn = re.search(r"\bОГРН[:\s]*(\d{13}|\d{15})\b", text, re.I) or re.search(
        r"\bОГРНИП[:\s]*(\d{15})\b", text, re.I
    )
    if ogrn:
        out["ogrn"] = ogrn.group(1)

    bik = re.search(r"\bБИК[:\s]*(\d{9})\b", text, re.I)
    if bik:
        out["bik"] = bik.group(1)

    acc = re.search(r"(?:р/?с|расч[её]тн(?:ый|ого)\s+сч[её]т)[:\s]*(\d{20})", text, re.I)
    if acc:
        out["account"] = acc.group(1)

    name = re.search(
        r"((?:ООО|ОАО|ПАО|АО|ИП|ЗАО)\s+[«\"]?[^«»\"\n,]+[»\"]?)",
        text,
        re.I,
    )
    if name:
        out["name"] = name.group(1).strip()
    elif not out.get("name"):
        out["name"] = "Контрагент без названия"

    return out


def save_counterparty(data: dict[str, str], raw_text: str = "") -> dict[str, Any]:
    init_db()
    inn = data.get("inn")
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        if inn:
            row = conn.execute(
                "SELECT * FROM counterparties WHERE inn = ?", (inn,)
            ).fetchone()
            if row:
                conn.execute(
                    """
                    UPDATE counterparties SET name=?, kpp=?, ogrn=?, account=?, bik=?, raw_text=?
                    WHERE inn=?
                    """,
                    (
                        data.get("name", row["name"]),
                        data.get("kpp") or row["kpp"],
                        data.get("ogrn") or row["ogrn"],
                        data.get("account") or row["account"],
                        data.get("bik") or row["bik"],
                        raw_text[:2000],
                        inn,
                    ),
                )
                conn.commit()
                return dict(conn.execute(
                    "SELECT * FROM counterparties WHERE inn = ?", (inn,)
                ).fetchone())

        conn.execute(
            """
            INSERT INTO counterparties (name, inn, kpp, ogrn, account, bik, raw_text, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                data.get("name", ""),
                data.get("inn"),
                data.get("kpp"),
                data.get("ogrn"),
                data.get("account"),
                data.get("bik"),
                raw_text[:2000],
                _now(),
            ),
        )
        conn.commit()
        rid = conn.execute("SELECT last_insert_rowid()").fetchone()[0]
        row = conn.execute("SELECT * FROM counterparties WHERE id = ?", (rid,)).fetchone()
        return dict(row) if row else {}


def list_counterparties(limit: int = 20) -> list[dict[str, Any]]:
    init_db()
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        rows = conn.execute(
            "SELECT id, name, inn, kpp, ogrn, account, bik, created_at FROM counterparties ORDER BY id DESC LIMIT ?",
            (limit,),
        ).fetchall()
        return [dict(r) for r in rows]


def _parse_1c_txt(content: bytes) -> list[dict[str, str]]:
    text = content.decode("utf-8", errors="replace")
    rows: list[dict[str, str]] = []
    for line in text.splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        parts = re.split(r"\t+|;", line)
        if len(parts) < 3:
            continue
        rows.append(
            {
                "date": parts[0].strip(),
                "counterparty": parts[1].strip() if len(parts) > 1 else "",
                "amount": parts[2].strip() if len(parts) > 2 else "",
                "purpose": parts[3].strip() if len(parts) > 3 else "",
            }
        )
    return rows


def _parse_xlsx(content: bytes) -> list[dict[str, str]]:
    try:
        import openpyxl
    except ImportError:
        return []

    wb = openpyxl.load_workbook(BytesIO(content), read_only=True, data_only=True)
    ws = wb.active
    headers: list[str] = []
    rows: list[dict[str, str]] = []
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        cells = [str(c).strip() if c is not None else "" for c in row]
        if i == 0:
            headers = [c.lower() for c in cells]
            continue
        if not any(cells):
            continue
        item: dict[str, str] = {}
        for h, v in zip(headers, cells):
            if "дата" in h or h == "date":
                item["date"] = v
            elif "контрагент" in h or "плательщик" in h or "получатель" in h:
                item["counterparty"] = v
            elif "сумм" in h or "amount" in h:
                item["amount"] = v
            elif "назнач" in h or "purpose" in h:
                item["purpose"] = v
        if not item and len(cells) >= 3:
            item = {
                "date": cells[0],
                "counterparty": cells[1],
                "amount": cells[2],
                "purpose": cells[3] if len(cells) > 3 else "",
            }
        if item:
            rows.append(item)
    return rows


def parse_bank_statement(filename: str, content: bytes) -> dict[str, Any]:
    """Парсинг выписки .xlsx или 1C_to_kl.txt."""
    ext = Path(filename).suffix.lower()
    transactions: list[dict[str, str]] = []

    if ext in (".xlsx", ".xls"):
        transactions = _parse_xlsx(content)
    elif ext in (".txt", ".csv"):
        transactions = _parse_1c_txt(content)
    else:
        return {"ok": False, "error": "Формат: .xlsx или .txt (1С)"}

    if not transactions:
        return {"ok": False, "error": "Транзакции не найдены — проверьте формат колонок"}

    income = 0.0
    expense = 0.0
    for t in transactions:
        raw = (t.get("amount") or "0").replace(" ", "").replace(",", ".")
        raw = re.sub(r"[^\d.\-]", "", raw)
        try:
            val = float(raw or 0)
        except ValueError:
            val = 0.0
        if val >= 0:
            income += val
        else:
            expense += abs(val)

    lines = [
        "| Дата | Контрагент | Сумма | Назначение |",
        "|------|------------|-------|------------|",
    ]
    for t in transactions[:30]:
        lines.append(
            f"| {t.get('date', '—')} | {t.get('counterparty', '—')[:24]} | "
            f"{t.get('amount', '—')} | {(t.get('purpose') or '—')[:40]} |"
        )
    if len(transactions) > 30:
        lines.append(f"\n_…и ещё {len(transactions) - 30} операций_")

    summary_md = (
        f"**Выписка:** `{filename}` — операций: **{len(transactions)}**\n"
        f"- Поступления (оценка): **{income:,.2f}** ₽\n"
        f"- Списания (оценка): **{expense:,.2f}** ₽\n\n"
        + "\n".join(lines)
    )

    global _last_statement_summary
    _last_statement_summary = summary_md

    path = UPLOADS_DIR / f"{uuid.uuid4().hex[:8]}_{Path(filename).name}"
    path.write_bytes(content)

    meta_path = path.with_suffix(path.suffix + ".json")
    meta_path.write_text(
        json.dumps({"transactions": transactions, "summary": summary_md}, ensure_ascii=False),
        encoding="utf-8",
    )

    return {
        "ok": True,
        "filename": filename,
        "stored": path.name,
        "transaction_count": len(transactions),
        "income_total": income,
        "expense_total": expense,
        "summary_markdown": summary_md,
    }


def get_statement_context() -> str:
    return _last_statement_summary


def get_counterparties_context() -> str:
    items = list_counterparties(10)
    if not items:
        return ""
    lines = ["[База контрагентов SQLite]"]
    for c in items:
        lines.append(
            f"- **{c.get('name') or '—'}** — ИНН {c.get('inn') or '—'}, "
            f"р/с {c.get('account') or '—'}, БИК {c.get('bik') or '—'}"
        )
    return "\n".join(lines)


def process_text_input(text: str) -> list[str]:
    """Обработка текста в режиме бухгалтера: реквизиты и т.д."""
    logs: list[str] = []
    req = parse_requisites(text)
    if req.get("inn") or req.get("account") or req.get("ogrn"):
        row = save_counterparty(req, raw_text=text)
        logs.append(
            f"[DB] Контрагент сохранён: {row.get('name')} (ИНН {row.get('inn') or '—'})"
        )
    return logs


def process_upload(filename: str, content: bytes) -> dict[str, Any]:
    ext = Path(filename).suffix.lower()
    if ext in (".xlsx", ".xls", ".txt", ".csv"):
        return parse_bank_statement(filename, content)
    return {
        "ok": True,
        "filename": filename,
        "indexed": True,
        "note": "Файл сохранён для RAG (общая загрузка)",
    }


def generate_contract_docx(counterparty_id: int | None = None) -> dict[str, str]:
    """Договор .docx (мок-шаблон, готов к расширению)."""
    init_db()
    cp: dict[str, Any] = {}
    if counterparty_id:
        with sqlite3.connect(DB_PATH) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute(
                "SELECT * FROM counterparties WHERE id = ?", (counterparty_id,)
            ).fetchone()
            if row:
                cp = dict(row)
    if not cp:
        cps = list_counterparties(1)
        cp = cps[0] if cps else {"name": "Контрагент", "inn": "—"}

    try:
        from docx import Document
    except ImportError:
        fid = uuid.uuid4().hex[:8]
        path = GENERATED_DIR / f"contract_{fid}.txt"
        path.write_text(
            f"ДОГОВОР ОКАЗАНИЯ УСЛУГ\n\nКонтрагент: {cp.get('name')}\nИНН: {cp.get('inn')}\n",
            encoding="utf-8",
        )
        return {
            "file_id": path.name,
            "filename": path.name,
            "download_url": f"/api/accountant/download/{path.name}",
            "format": "txt",
            "message": "Установите python-docx для .docx",
        }

    doc = Document()
    doc.add_heading("Договор оказания услуг", 0)
    doc.add_paragraph(f"Контрагент: {cp.get('name', '—')}")
    doc.add_paragraph(f"ИНН: {cp.get('inn', '—')}, КПП: {cp.get('kpp', '—')}")
    doc.add_paragraph(f"ОГРН: {cp.get('ogrn', '—')}")
    doc.add_paragraph(f"р/с: {cp.get('account', '—')}, БИК: {cp.get('bik', '—')}")
    doc.add_paragraph(
        "Настоящий договор составлен в соответствии с законодательством РФ (ГК РФ). "
        "Шаблон сгенерирован Jarvis — уточните условия у юриста."
    )

    fname = f"contract_{uuid.uuid4().hex[:8]}.docx"
    path = GENERATED_DIR / fname
    doc.save(path)
    return {
        "file_id": fname,
        "filename": fname,
        "download_url": f"/api/accountant/download/{fname}",
        "format": "docx",
        "message": "Договор сформирован",
    }


def generate_invoice_xlsx(counterparty_id: int | None = None, amount: float = 10000.0) -> dict[str, str]:
    """Счёт .xlsx (мок-шаблон)."""
    init_db()
    cp: dict[str, Any] = {}
    if counterparty_id:
        with sqlite3.connect(DB_PATH) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute(
                "SELECT * FROM counterparties WHERE id = ?", (counterparty_id,)
            ).fetchone()
            if row:
                cp = dict(row)
    if not cp:
        cps = list_counterparties(1)
        cp = cps[0] if cps else {"name": "Контрагент", "inn": "—"}

    try:
        from openpyxl import Workbook
    except ImportError:
        fid = uuid.uuid4().hex[:8]
        path = GENERATED_DIR / f"invoice_{fid}.csv"
        path.write_text(
            f"Счёт;{cp.get('name')};{amount}\nИНН;{cp.get('inn')}\n",
            encoding="utf-8",
        )
        return {
            "file_id": path.name,
            "filename": path.name,
            "download_url": f"/api/accountant/download/{path.name}",
            "format": "csv",
            "message": "Установите openpyxl для .xlsx",
        }

    wb = Workbook()
    ws = wb.active
    ws.title = "Счёт"
    ws.append(["Счёт на оплату"])
    ws.append(["Покупатель", cp.get("name", "")])
    ws.append(["ИНН", cp.get("inn", "")])
    ws.append(["Сумма", amount])
    ws.append(["НДС", "Без НДС / уточните"])
    fname = f"invoice_{uuid.uuid4().hex[:8]}.xlsx"
    path = GENERATED_DIR / fname
    wb.save(path)
    return {
        "file_id": fname,
        "filename": fname,
        "download_url": f"/api/accountant/download/{fname}",
        "format": "xlsx",
        "message": "Счёт сформирован",
    }


def get_generated_path(file_id: str) -> Path | None:
    safe = Path(file_id).name
    path = GENERATED_DIR / safe
    if path.is_file() and path.parent.resolve() == GENERATED_DIR.resolve():
        return path
    return None


def build_accountant_context_extra() -> str:
    parts = []
    cp = get_counterparties_context()
    st = get_statement_context()
    if cp:
        parts.append(cp)
    if st:
        parts.append(f"\n[Последняя банковская выписка]\n{st}")
    return "\n\n".join(parts)
